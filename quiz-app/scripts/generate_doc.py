#!/usr/bin/env python3
"""Generate Word (.docx) or PDF files from content."""
import sys
import json
import base64
import re
import os
import io

# Fix Windows stdin/stdout encoding for Thai/Unicode
if sys.platform == 'win32':
    sys.stdin = io.TextIOWrapper(sys.stdin.buffer, encoding='utf-8', errors='replace')
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

def strip_markdown(text):
    blocks = []
    lines = text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.startswith('### '):
            blocks.append({'type': 'h3', 'text': line[4:].strip()})
        elif line.startswith('## '):
            blocks.append({'type': 'h2', 'text': line[3:].strip()})
        elif line.startswith('# '):
            blocks.append({'type': 'h1', 'text': line[2:].strip()})
        elif line.startswith('**') and line.endswith('**'):
            blocks.append({'type': 'bold', 'text': line.strip('* ')})
        elif line.startswith('- ') or line.startswith('* '):
            blocks.append({'type': 'bullet', 'text': clean_inline(line[2:].strip())})
        elif re.match(r'^\d+\.\s', line):
            blocks.append({'type': 'numbered', 'text': clean_inline(re.sub(r'^\d+\.\s', '', line).strip())})
        elif line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i])
                i += 1
            blocks.append({'type': 'code', 'text': '\n'.join(code_lines)})
        elif '|' in line and line.strip().startswith('|'):
            table_rows = []
            while i < len(lines) and '|' in lines[i] and lines[i].strip().startswith('|'):
                cells = [c.strip() for c in lines[i].strip('| ').split('|')]
                if not all(re.match(r'^-+:?$|^:?-+:?$', c) for c in cells):
                    table_rows.append(cells)
                i += 1
            i -= 1
            blocks.append({'type': 'table', 'rows': table_rows})
        elif line.strip() == '':
            blocks.append({'type': 'empty'})
        else:
            blocks.append({'type': 'paragraph', 'text': clean_inline(line.strip())})
        i += 1
    return blocks

def clean_inline(text):
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    text = re.sub(r'`([^`]+)`', r'\1', text)
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
    return text

def add_rich_text(paragraph, text, font_size=14, font_name='TH Sarabun New'):
    """Add text with inline bold/italic/code formatting to a paragraph."""
    from docx.shared import Pt, RGBColor
    # Split by **bold** and *italic* patterns
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`[^`]+`)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.name = font_name
            run.font.size = Pt(font_size)
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
            run.font.name = font_name
            run.font.size = Pt(font_size)
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        else:
            # Clean link syntax [text](url) -> text
            cleaned = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', part)
            run = paragraph.add_run(cleaned)
            run.font.name = font_name
            run.font.size = Pt(font_size)

def generate_docx(content, title):
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'TH Sarabun New'
    font.size = Pt(14)

    if title:
        p = doc.add_heading(title, level=0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    blocks = strip_markdown(content)
    for block in blocks:
        btype = block['type']
        text = block.get('text', '')
        if btype == 'h1':
            doc.add_heading(clean_inline(text), level=1)
        elif btype == 'h2':
            doc.add_heading(clean_inline(text), level=2)
        elif btype == 'h3':
            doc.add_heading(clean_inline(text), level=3)
        elif btype == 'bold':
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(14)
            run.font.name = 'TH Sarabun New'
        elif btype == 'bullet':
            p = doc.add_paragraph(style='List Bullet')
            add_rich_text(p, text)
        elif btype == 'numbered':
            p = doc.add_paragraph(style='List Number')
            add_rich_text(p, text)
        elif btype == 'code':
            p = doc.add_paragraph()
            run = p.add_run(block['text'])
            run.font.name = 'Consolas'
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            p.paragraph_format.left_indent = Inches(0.5)
        elif btype == 'table':
            rows = block.get('rows', [])
            if rows:
                col_count = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=col_count)
                table.style = 'Table Grid'
                for r, row in enumerate(rows):
                    for c, cell in enumerate(row):
                        if c < col_count:
                            table.cell(r, c).text = clean_inline(cell)
                            if r == 0:
                                for paragraph in table.cell(r, c).paragraphs:
                                    for run in paragraph.runs:
                                        run.bold = True
        elif btype == 'paragraph' and text:
            p = doc.add_paragraph()
            add_rich_text(p, text)

    import io
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

def generate_pdf(content, title):
    from fpdf import FPDF
    import urllib.request
    import tempfile

    font_dir = os.path.join(tempfile.gettempdir(), 'thai_fonts')
    os.makedirs(font_dir, exist_ok=True)
    font_path = os.path.join(font_dir, 'Sarabun-Regular.ttf')
    font_bold_path = os.path.join(font_dir, 'Sarabun-Bold.ttf')

    for fp, url in [
        (font_path, 'https://github.com/google/fonts/raw/main/ofl/sarabun/Sarabun-Regular.ttf'),
        (font_bold_path, 'https://github.com/google/fonts/raw/main/ofl/sarabun/Sarabun-Bold.ttf'),
    ]:
        if not os.path.exists(fp):
            try:
                urllib.request.urlretrieve(url, fp)
            except:
                pass

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    has_thai = False
    if os.path.exists(font_path):
        try:
            pdf.add_font('Sarabun', '', font_path)
            if os.path.exists(font_bold_path):
                pdf.add_font('Sarabun', 'B', font_bold_path)
            has_thai = True
        except:
            pass

    fn = 'Sarabun' if has_thai else 'Helvetica'
    usable_w = pdf.w - pdf.l_margin - pdf.r_margin

    def nl():
        pdf.set_x(pdf.l_margin)

    def write_text(text, size=14, bold=False, indent=0):
        nl()
        if indent:
            pdf.set_x(pdf.l_margin + indent)
        pdf.set_font(fn, 'B' if bold else '', size)
        w = pdf.w - pdf.r_margin - pdf.get_x()
        if w < 20:
            nl()
            w = usable_w
        pdf.multi_cell(w, 7, text)

    if title:
        nl()
        pdf.set_font(fn, 'B', 20)
        pdf.cell(usable_w, 12, title, new_x="LMARGIN", new_y="NEXT", align='C')
        pdf.ln(5)

    blocks = strip_markdown(content)

    for block in blocks:
        btype = block['type']
        if btype == 'h1':
            write_text(block['text'], size=18, bold=True)
            pdf.ln(2)
        elif btype == 'h2':
            write_text(block['text'], size=16, bold=True)
            pdf.ln(2)
        elif btype == 'h3':
            write_text(block['text'], size=14, bold=True)
            pdf.ln(1)
        elif btype == 'bold':
            write_text(block['text'], bold=True)
            pdf.ln(1)
        elif btype == 'bullet':
            write_text('  -  ' + block['text'], indent=5)
        elif btype == 'numbered':
            write_text('  ' + block['text'], indent=5)
        elif btype == 'code':
            code_font = fn if has_thai else 'Courier'
            pdf.set_font(code_font, '', 10)
            pdf.set_fill_color(245, 245, 245)
            for code_line in block['text'].split('\n'):
                nl()
                safe = '  ' + code_line[:100]
                pdf.cell(usable_w, 5, safe, new_x="LMARGIN", new_y="NEXT", fill=True)
            pdf.ln(2)
        elif btype == 'table':
            rows = block.get('rows', [])
            if rows:
                col_count = max(len(r) for r in rows) if rows else 1
                col_w = usable_w / max(col_count, 1)
                for r, row in enumerate(rows):
                    nl()
                    pdf.set_font(fn, 'B' if r == 0 else '', 11)
                    pdf.set_fill_color(230 if r == 0 else 255, 230 if r == 0 else 255, 230 if r == 0 else 255)
                    for c in range(col_count):
                        txt = clean_inline(row[c])[:30] if c < len(row) else ''
                        pdf.cell(col_w, 8, txt, border=1, fill=(r == 0))
                    pdf.cell(0, 0, '', new_x="LMARGIN", new_y="NEXT")
                pdf.ln(3)
        elif btype == 'paragraph' and block.get('text'):
            write_text(block['text'])
            pdf.ln(1)
        elif btype == 'empty':
            pdf.ln(3)

    return pdf.output()

def main():
    input_data = json.loads(sys.stdin.read())
    content = input_data.get('content', '')
    title = input_data.get('title', '')
    fmt = input_data.get('format', 'docx')

    try:
        if fmt == 'docx':
            result = generate_docx(content, title)
        elif fmt == 'pdf':
            result = generate_pdf(content, title)
        else:
            print(json.dumps({'error': f'Unknown format: {fmt}'}))
            sys.exit(1)

        encoded = base64.b64encode(result).decode('utf-8')
        print(json.dumps({'data': encoded, 'format': fmt}))
    except Exception as e:
        print(json.dumps({'error': str(e)}))
        sys.exit(1)

if __name__ == '__main__':
    main()
